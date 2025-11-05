#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为Word文档（改进版）
"""
import argparse
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown_table(lines: list, start_idx: int) -> tuple:
    """解析Markdown表格，返回(表格数据, 结束索引)"""
    table_data = []
    i = start_idx
    
    # 跳过分隔行（|---|---|---|）
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith('|'):
            break
        if '---' in line:
            i += 1
            continue
        
        # 解析表格行
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        if cells:
            table_data.append(cells)
        i += 1
    
    return table_data, i


def add_table_to_doc(doc: Document, table_data: list):
    """添加表格到Word文档"""
    if not table_data or len(table_data) < 2:
        return
    
    # 创建表格
    num_rows = len(table_data)
    num_cols = max(len(row) for row in table_data) if table_data else 0
    
    if num_cols == 0:
        return
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # 填充表格
    for i, row_data in enumerate(table_data):
        for j in range(num_cols):
            cell = table.rows[i].cells[j]
            if j < len(row_data):
                cell.text = row_data[j]
            else:
                cell.text = ""
            
            # 设置单元格格式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    if hasattr(run._element.rPr, 'rFonts'):
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def process_text_with_formatting(text: str, paragraph):
    """处理带格式的文本（加粗、列表等）"""
    # 处理加粗 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            if hasattr(run._element.rPr, 'rFonts'):
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif part:
            run = paragraph.add_run(part)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            if hasattr(run._element.rPr, 'rFonts'):
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def markdown_to_word(md_file: str, output_file: str):
    """将Markdown文件转换为Word文档"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 标题处理
        if line.strip().startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif level == 2:
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif level == 3:
                p = doc.add_heading(text, level=3)
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            else:
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.size = Pt(13)
                    run.font.bold = True
        
        # 表格处理
        elif line.strip().startswith('|'):
            table_data, next_idx = parse_markdown_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
            i = next_idx
            continue
        
        # 列表项处理
        elif line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.name = '宋体'
                if hasattr(run._element.rPr, 'rFonts'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 普通段落处理
        else:
            text = line.strip()
            # 跳过只有标记符号的行（如分隔线）
            if text in ['---', '***', '___']:
                i += 1
                continue
            
            p = doc.add_paragraph()
            process_text_with_formatting(text, p)
            
            # 设置段落格式
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
        
        i += 1
    
    # 保存文档
    doc.save(output_file)
    print(f"已生成Word文档: {output_file}")


def main():
    parser = argparse.ArgumentParser(description="将Markdown文件转换为Word文档")
    parser.add_argument("input", help="输入的Markdown文件路径")
    parser.add_argument("-o", "--output", help="输出的Word文档路径（可选）")
    args = parser.parse_args()
    
    # 确定输出文件路径
    if args.output:
        output_file = args.output
    else:
        base_name = os.path.splitext(args.input)[0]
        output_file = f"{base_name}.docx"
    
    markdown_to_word(args.input, output_file)


if __name__ == "__main__":
    main()

