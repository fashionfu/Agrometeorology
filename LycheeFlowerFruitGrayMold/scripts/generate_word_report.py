#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将Markdown格式的预警等级阈值规则报告转换为Word文档
"""

import os
import re
import argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown_to_word(md_file: str, output_file: str):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行（后面统一处理）
        if not line:
            i += 1
            continue
        
        # 一级标题
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        
        # 二级标题
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
        
        # 三级标题
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
        
        # 四级标题
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
            i += 1
        
        # 分隔线
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
            i += 1
        
        # 粗体文本（**text**） - 需要检查是否在表格或代码块中
        elif '**' in line and not line.startswith('|') and not line.startswith('```'):
            para = doc.add_paragraph()
            # 处理粗体文本（支持多个粗体片段）
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    para.add_run(part)
            i += 1
        
        # 表格行检测（包含 | 的行）
        elif '|' in line and line.count('|') >= 3:
            # 收集表格行
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            # 解析表格
            if table_lines:
                # 跳过分隔行（只包含 |--- 的行）
                header_line = table_lines[0]
                data_lines = [l for l in table_lines[1:] if not re.match(r'^\|[\s\-\|:]+\|$', l)]
                
                if data_lines:
                    # 解析表头
                    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
                    
                    # 创建表格
                    table = doc.add_table(rows=1, cols=len(headers))
                    # 尝试使用网格样式，如果不存在则使用默认样式
                    try:
                        table.style = 'Light Grid Accent 1'
                    except:
                        try:
                            table.style = 'Table Grid'
                        except:
                            pass  # 使用默认样式
                    
                    # 添加表头
                    header_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        header_cells[j].text = header
                        header_cells[j].paragraphs[0].runs[0].bold = True
                    
                    # 添加数据行
                    for data_line in data_lines:
                        cells = [cell.strip() for cell in data_line.split('|')[1:-1]]
                        row = table.add_row()
                        for j, cell in enumerate(cells):
                            row.cells[j].text = cell
            continue
        
        # 代码块（``` 包围的内容）
        elif line.startswith('```'):
            # 跳过代码块标记
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # 跳过结束标记
            if i < len(lines):
                i += 1
            
            # 添加代码块（使用等宽字体）
            if code_lines:
                para = doc.add_paragraph()
                para.style = 'No Spacing'
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 128)
        
        # 列表项（以 - 或 * 开头）
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            para = doc.add_paragraph(text, style='List Bullet')
            i += 1
        
        # 数字列表
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            para = doc.add_paragraph(text, style='List Number')
            i += 1
        
        # 普通段落
        else:
            para = doc.add_paragraph(line)
            i += 1
        
        # 在段落后添加小间距
        if i < len(lines):
            # 检查下一行是否为空，如果不是特定结构，添加小间距
            next_line = lines[i].strip() if i < len(lines) else ''
            if next_line and not next_line.startswith('#') and not next_line.startswith('|'):
                pass  # 不添加额外间距
    
    # 保存文档
    doc.save(output_file)
    print(f"✓ Word文档已生成: {output_file}")


def main():
    parser = argparse.ArgumentParser(description="将Markdown报告转换为Word文档")
    parser.add_argument(
        "--input",
        default="analysis_1105/预警等级阈值规则报告.md",
        help="输入的Markdown文件路径"
    )
    parser.add_argument(
        "--output",
        default="analysis_1105/预警等级阈值规则报告.docx",
        help="输出的Word文档路径"
    )
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"错误: 找不到输入文件: {args.input}")
        return
    
    try:
        parse_markdown_to_word(args.input, args.output)
        print("转换完成！")
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
