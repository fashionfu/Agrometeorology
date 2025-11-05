#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为Word文档
"""
import argparse
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown_table(line: str) -> list:
    """解析Markdown表格行"""
    # 移除首尾的|符号
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    
    # 分割单元格
    cells = [cell.strip() for cell in line.split('|')]
    return cells


def add_table_to_doc(doc: Document, lines: list, start_idx: int) -> int:
    """添加表格到文档，返回处理后的行索引"""
    # 找到表格结束位置
    end_idx = start_idx
    while end_idx < len(lines) and (lines[end_idx].strip().startswith('|') or lines[end_idx].strip() == ''):
        end_idx += 1
    
    # 提取表格数据
    table_data = []
    for i in range(start_idx, end_idx):
        line = lines[i].strip()
        if line.startswith('|') and '---' not in line:
            cells = parse_markdown_table(line)
            if cells:
                table_data.append(cells)
    
    if len(table_data) < 2:  # 至少需要表头和一行数据
        return end_idx
    
    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    
    # 填充表格
    for i, row_data in enumerate(table_data):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                # 设置单元格格式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    return end_idx


def markdown_to_word(md_file: str, output_file: str):
    """将Markdown文件转换为Word文档"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题处理
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                # 一级标题
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif level == 2:
                # 二级标题
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif level == 3:
                # 三级标题
                p = doc.add_heading(text, level=3)
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            else:
                # 其他级别标题
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.size = Pt(13)
                    run.font.bold = True
        
        # 表格处理
        elif line.startswith('|'):
            i = add_table_to_doc(doc, lines, i)
            continue
        
        # 普通段落处理
        else:
            # 处理加粗文本
            text = line
            p = doc.add_paragraph()
            
            # 处理加粗标记
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # 加粗文本
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(12)
                elif part:
                    # 普通文本
                    run = p.add_run(part)
                    run.font.size = Pt(12)
            
            # 设置段落格式
            p.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
            p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        
        i += 1
    
    # 保存文档
    doc.save(output_file)
    print(f"已生成Word文档: {output_file}")


def main():
    parser = argparse.ArgumentParser(description="将Markdown文件转换为Word文档")
    parser.add_argument("input", help="输入的Markdown文件路径")
    parser.add_argument("-o", "--output", help="输出的Word文档路径（可选）")
    args = parser.parse_args()
    
    # 确定输出文件路径
    if args.output:
        output_file = args.output
    else:
        base_name = os.path.splitext(args.input)[0]
        output_file = f"{base_name}.docx"
    
    markdown_to_word(args.input, output_file)


if __name__ == "__main__":
    main()


