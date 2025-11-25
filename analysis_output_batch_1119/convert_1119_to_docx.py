#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将 1119分析.md 转换为 Word 文档
"""
import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def markdown_to_docx(md_file, docx_file):
    """将 Markdown 文件转换为 Word 文档"""
    doc = Document()
    
    # 设置页面大小（A4）
    doc.sections[0].page_width = Inches(8.27)
    doc.sections[0].page_height = Inches(11.69)
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)
    doc.sections[0].top_margin = Inches(1)
    doc.sections[0].bottom_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题处理
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        # 表格处理
        elif line.startswith('|'):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                # 跳过分隔行
                if not all(cell.strip() in ['', '-', ':', '--'] for cell in row):
                    table_data.append(row)
                i += 1
            i -= 1
            
            if table_data and len(table_data) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table_data[0]):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # 设置单元格字体
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                                    run.font.name = '宋体'
                            # 表头加粗
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
        # 代码块
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        # 列表项
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
        # 普通段落（处理粗体）
        else:
            p = doc.add_paragraph()
            # 处理粗体文本 **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        i += 1
    
    doc.save(docx_file)
    print(f"文档已保存到: {docx_file}")

if __name__ == '__main__':
    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, '1119分析.md')
    docx_file = os.path.join(script_dir, '1119分析.docx')
    
    if not os.path.exists(md_file):
        print(f"错误: 找不到文件 {md_file}")
        sys.exit(1)
    
    print("正在转换 Markdown 文档为 Word 文档...")
    markdown_to_docx(md_file, docx_file)
    print("转换完成！")


